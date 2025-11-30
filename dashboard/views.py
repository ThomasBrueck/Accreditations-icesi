import csv
from django.shortcuts import redirect, render, get_object_or_404, HttpResponseRedirect
from django.views import View
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from django.http import JsonResponse
from django.core.paginator import Paginator
from django.contrib.auth import get_user_model
from django.views.generic import TemplateView, UpdateView, DeleteView
from django.views.generic.edit import CreateView
from django.urls import reverse_lazy, reverse
from django.views.generic import ListView
from init.models import RoleModel, UserModel
from django.contrib import messages
from .models import CharacteristicAspects, CharacteristicModel, CharacteristicStrengths, GlobalAspects, GlobalStrengths, QuestionModel, ReportModel, FactorModel, CommentsModel, NotificationModel, States
from .forms import AnswerForm, CharacteristicDevelopForm, CharacteristicForm, QuestionForm, ReportFilterForm, ReportForm, FactorForm
from django.utils import timezone
from datetime import datetime, timedelta
from django.core.exceptions import PermissionDenied
from .forms import ProfileForm, TaskForm
from django.contrib import messages
from django.shortcuts import render, redirect
from django.utils.dateparse import parse_date
from django.contrib import messages
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required, permission_required
from django.core.mail import send_mail
from django.conf import settings
from django.contrib.auth.models import User
from .models import NotificationLog
from django.core.validators import validate_email
from django.core.exceptions import ValidationError
from .models import AccreditationProcess
from django.http import HttpResponse
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
# Create your views here.

@method_decorator(login_required, name="dispatch")
class UserDashboardView(View):
    def get(self, request):
        # Basic statistics
        total_reports = ReportModel.objects.count()

        recent_comments = CommentsModel.objects.select_related('owner', 'factor').order_by('-created_at')[:5]
        recent_reports = ReportModel.objects.select_related('created_by').order_by('-created_at')[:5]

        today = timezone.now()
        thirty_days_ago = today - timedelta(days=30)
        
        active_reports = ReportModel.objects.filter(status='active').count()
        inactive_reports = ReportModel.objects.filter(status='inactive').count()
        
        recent_created_reports = ReportModel.objects.filter(created_at__gte=thirty_days_ago).count()
        
        factors_with_content = FactorModel.objects.exclude(content='').count()
        total_factors = FactorModel.objects.count()
        completion_rate = int((factors_with_content / total_factors * 100) if total_factors > 0 else 0)

        context = {
            "total_reports": ReportModel.objects.count(),
            "recent_comments": CommentsModel.objects.order_by("-created_at")[:3],
            "recent_reports": ReportModel.objects.order_by("-created_at")[:3],
            "active_count": ReportModel.objects.filter(status='active').count(),
            "inactive_count": ReportModel.objects.filter(status='inactive').count(),
        }

        return render(request, 'dashboard/dashboard-user.html', context)
    
@method_decorator(login_required, name="dispatch")
class AdminDashboardView(View):
    def get(self, request):
        # Basic statistics
        total_reports = ReportModel.objects.count()
        total_users_assigned = UserModel.objects.filter(is_superuser=False).exclude(role__name='common').count()
        pending_approvals = CommentsModel.objects.filter(status='pending').count()

        # Recent activity
        recent_comments = CommentsModel.objects.select_related('owner', 'factor').order_by('-created_at')[:5]
        recent_reports = ReportModel.objects.select_related('created_by').order_by('-created_at')[:5]
        recent_users = UserModel.objects.filter(is_superuser=False).order_by('-date_joined')[:5]

        # Monthly statistics
        today = timezone.now()
        thirty_days_ago = today - timedelta(days=30)
        
        # Reports by status
        active_reports = ReportModel.objects.filter(status='active').count()
        inactive_reports = ReportModel.objects.filter(status='inactive').count()
        
        # Reports created in the last 30 days
        recent_created_reports = ReportModel.objects.filter(created_at__gte=thirty_days_ago).count()
        
        # Calculate completion rate
        factors_with_content = FactorModel.objects.exclude(content='').count()
        total_factors = FactorModel.objects.count()
        completion_rate = int((factors_with_content / total_factors * 100) if total_factors > 0 else 0)

        context = {
            "total_reports": ReportModel.objects.count(),
            "total_users_assigned": UserModel.objects.exclude(is_superuser=True).count(),
            "pending_approvals": CommentsModel.objects.filter(status="pending").count(),
            "recent_comments": CommentsModel.objects.order_by("-created_at")[:3],
            "recent_reports": ReportModel.objects.order_by("-created_at")[:3],
            "active_count": ReportModel.objects.filter(status='active').count(),
            "inactive_count": ReportModel.objects.filter(status='inactive').count(),
        }

        return render(request, 'dashboard/dashboard-admin.html', context)


@method_decorator(login_required, name="dispatch")
class MarkNotificationsReadView(View):
    def post(self, request, *args, **kwargs):
        if request.user.is_authenticated:
            NotificationModel.objects.filter(user=request.user, is_read=False).update(is_read=True)
            return JsonResponse({'status': 'success'})
        raise PermissionDenied


@method_decorator(login_required(login_url='login-user'), name='dispatch')
class AssignRoleView(View):
    def get(self, request):
        users = UserModel.objects.all()
        roles = RoleModel.objects.all()
        return render(request, "dashboard/assign_role.html", {
            "users": users,
            "roles": roles,
        })

    def post(self, request):
        user_role = getattr(request.user, 'role', None)
        if not user_role or user_role.name not in ['acadi', 'program director']:
            messages.error(request, 
                "Solo los usuarios con rol “acadi” o “program director” pueden asignar roles."
            )
            return redirect('assign-role-view')
        
        user_id = request.POST.get("user_id")
        role_id = request.POST.get("role_id")
        if not user_id or not role_id:
            messages.error(request, "Debes seleccionar un usuario y un rol.")
            return redirect('assign-role-view')

        try:
            target_user = UserModel.objects.get(pk=int(user_id))
            new_role    = RoleModel.objects.get(pk=int(role_id))
        except (UserModel.DoesNotExist, RoleModel.DoesNotExist, ValueError):
            messages.error(request, "Usuario o rol no válidos.")
            return redirect('assign-role-view')

        target_user.role = new_role
        target_user.save()

        messages.success(request,
            f"Rol “{new_role.name}” asignado con éxito a {target_user.username}."
        )
        return redirect('assign-role-view')
    

@method_decorator(login_required(login_url='login-user'), name='dispatch')
class RemoveRoleView(View):
    def dispatch(self, request, *args, **kwargs):
        role = getattr(request.user, 'role', None)
        if not role or role.name.lower() != 'acadi':
            return redirect('dashboard-admin')
        return super().dispatch(request, *args, **kwargs)

    def get(self, request):
        users = UserModel.objects.all()
        return render(request, 'dashboard/remove_role.html', {'users': users})

    def post(self, request):
        user_id = request.POST.get('user_id')
        try:
            target = UserModel.objects.get(pk=int(user_id))
        except (UserModel.DoesNotExist, ValueError):
            messages.error(request, "Usuario no válido.")
            return redirect('remove-role-view')

        common = RoleModel.objects.get(name='common')
        target.role = common
        target.save()

        messages.success(request, f"Rol de {target.username} cambiado a 'common' con éxito.")
        return redirect('remove-role-view')
    

@method_decorator(login_required, name="dispatch")
class ReportListView(ListView):
    model = ReportModel
    template_name = 'dashboard/reports.html'
    context_object_name = 'reports'
    paginate_by = 10
    
    def get_queryset(self):
        queryset = super().get_queryset()
        search_query = self.request.GET.get('search', '')
        filter_form = ReportFilterForm(self.request.GET)

        if filter_form.is_valid():
            name = filter_form.cleaned_data.get('name')
            status = filter_form.cleaned_data.get('status')
            created_by = filter_form.cleaned_data.get('created_by')
            end_date = filter_form.cleaned_data.get('end_date')

            if search_query:
                queryset = queryset.filter(name__icontains=search_query)

            if name:
                queryset = queryset.filter(name__icontains=name)
            if status and status != '':
                queryset = queryset.filter(status=status)
            if created_by:
                queryset = queryset.filter(created_by__username__icontains=created_by)
            if end_date:
                queryset = queryset.filter(end_date__lte=end_date)

        return queryset
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['search_query'] = self.request.GET.get('search', '')
        context['filter_form'] = ReportFilterForm(self.request.GET)
        end_date_str = self.request.GET.get('end_date', '')
        end_date = parse_date(end_date_str) if end_date_str else None
        context['active_filters'] = {
            'name': self.request.GET.get('name', ''),
            'status': self.request.GET.get('status', ''),
            'created_by': self.request.GET.get('created_by', ''),
            'end_date': end_date  # Pasamos un objeto date para que el template lo maneje
        }
        context['user'] = self.request.user
        return context


@method_decorator(login_required, name="dispatch")
class ReportCreateView(CreateView):
    model = ReportModel
    form_class = ReportForm
    template_name = 'dashboard/report-form.html'
    success_url = reverse_lazy('report-list')
    
    def form_valid(self, form):
        form.instance.created_by = self.request.user
        response = super().form_valid(form)

        for i in range(12):
            FactorModel.objects.create(
                name=f"Factor {i+1}",
                content="",
                report=self.object,
                last_edited_by=self.request.user
            )

        messages.success(self.request, f"Report '{self.object.name}' with 12 factors created successfully")
        return response


@method_decorator(login_required, name="dispatch")
class ReportUpdateView(UpdateView):
    model = ReportModel
    form_class = ReportForm
    template_name = 'dashboard/report-form.html'
    success_url = reverse_lazy('report-list')
    
    def form_valid(self, form):
        response = super().form_valid(form)
        messages.success(self.request, f"Report '{self.object.name}' updated succesfully")
        return response


#Necesario porque no se quiere renderizar una template para delete report
@login_required
def ReportDeleteView(request, pk):
    report = get_object_or_404(ReportModel, pk=pk)
    
    if request.method == 'POST':
        name = report.name 
        report.delete()
        messages.success(request, f"Report '{name}' deleted succesfully")
        return redirect('report-list')
    
    return redirect('report-list')


@method_decorator(login_required, name="dispatch")
class ReportFactorView(ListView):
    model = FactorModel
    template_name = 'dashboard/report-detail.html'
    context_object_name = 'factors'
    paginate_by = 6

    def get_queryset(self):
        self.report = get_object_or_404(ReportModel, pk=self.kwargs['pk'])
        return self.report.factors.all().order_by('id')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['report'] = self.report
        return context

@method_decorator(login_required, name="dispatch")
class CommentsListView(View):
    def get(self, request, factor_id):
        factor = get_object_or_404(FactorModel, id=factor_id)
        comments = factor.comments.all()
        return render(request, 'dashboard/comments_list.html', {
            'factor': factor,
            'comments': comments
        })

@method_decorator(login_required, name="dispatch")
class CommentCreateView(View):
    def get(self, request, factor_id):
        factor = get_object_or_404(FactorModel, id=factor_id)
        return render(request, 'dashboard/comment_create.html', {'factor': factor})

    def post(self, request, factor_id):
        factor = get_object_or_404(FactorModel, id=factor_id)
        title = request.POST.get('title')
        content = request.POST.get('content')

        if title and content:
            CommentsModel.objects.create(
                factor=factor,
                owner=request.user,
                title=title,
                content=content
                # No hace falta pasar status, se asigna 'pending' por defecto
            )
            messages.success(request, "Comentario creado exitosamente.")
            return redirect('comments-list', factor_id=factor.id)
        else:
            messages.error(request, "Por favor, completa todos los campos.")
            return render(request, 'dashboard/comment_create.html', {'factor': factor})
        
@method_decorator(login_required, name="dispatch")
class CommentUpdateView(View):
    def get(self, request, comment_id):
        comment = get_object_or_404(CommentsModel, id=comment_id)
        # Verificar que el usuario autenticado sea el dueño del comentario
        if comment.owner != request.user:
            messages.error(request, "No tienes permiso para editar este comentario.")
            return redirect('comments-list', factor_id=comment.factor.id)
        return render(request, 'dashboard/comment_create.html', {
            'factor': comment.factor,
            'comment': comment
        })

    def post(self, request, comment_id):
        comment = get_object_or_404(CommentsModel, id=comment_id)
        # Verificar que el usuario autenticado sea el dueño del comentario
        if comment.owner != request.user:
            messages.error(request, "No tienes permiso para editar este comentario.")
            return redirect('comments-list', factor_id=comment.factor.id)

        title = request.POST.get('title')
        content = request.POST.get('content')
        status = request.POST.get('status')

        if title and content:
            comment.title = title
            comment.content = content
            comment.save()
            messages.success(request, "Comment update succesfully.")
            return redirect('comments-list', factor_id=comment.factor.id)
        else:
            messages.error(request, "Please, complete all the fields.")
            return render(request, 'dashboard/comment_create.html', {
                'factor': comment.factor,
                'comment': comment
            })

@method_decorator(login_required, name="dispatch")
class CommentDeleteView(View):
    def post(self, request, comment_id):
        comment = get_object_or_404(CommentsModel, id=comment_id)
        if comment.owner != request.user:
            messages.error(request, "No tienes permiso para eliminar este comentario.")
            return redirect('comments-list', factor_id=comment.factor.id)

        factor_id = comment.factor.id
        comment.delete()
        messages.success(request, "Comment delete succesfully.")
        return redirect('comments-list', factor_id=factor_id)
        
@method_decorator(login_required, name="dispatch")
class CommentDetailView(View):
    def get(self, request, comment_id):
        comment = get_object_or_404(CommentsModel, id=comment_id)
        return render(request, 'dashboard/comment_detail.html', {
            'comment': comment
        })
    
@method_decorator(login_required, name="dispatch")
class CommentsReviewListView(View):
    def get(self, request, factor_id):
        factor = get_object_or_404(FactorModel, id=factor_id)
        user_role = getattr(request.user, 'role', None)
        if not user_role or user_role.name not in ['acadi', 'program director']:
            messages.error(request, "Only users with role 'acadi' or 'program director' can review comments.")
            return redirect('comments-list', factor_id=factor.id)

        comments = factor.comments.all()
        return render(request, 'dashboard/comments_list.html', {
            'factor': factor,
            'comments': comments
        })

@method_decorator(login_required, name="dispatch")
class CommentReviewView(View):
    def dispatch(self, request, *args, **kwargs):
        user_role = getattr(request.user, 'role', None)
        if not user_role or user_role.name not in ['acadi', 'program director']:
            messages.error(request, "Only users with role 'acadi' or 'program director' can review comments.")
            return redirect('dashboard-admin')
        return super().dispatch(request, *args, **kwargs)

    def get(self, request, comment_id):
        comment = get_object_or_404(CommentsModel, id=comment_id)
        return render(request, 'dashboard/comment_review.html', {
            'comment': comment
        })

    def post(self, request, comment_id):
        comment = get_object_or_404(CommentsModel, id=comment_id)
        action = request.POST.get('action')
        justification = request.POST.get('justification', '').strip()

        if action == 'approve':
            comment.status = 'approved'
            if justification:
                comment.justification = justification
            messages.success(request, "Comment approved successfully.")
        elif action == 'disapprove':
            if not justification:
                messages.error(request, "Justification is required when disapproving a comment.")
                return render(request, 'dashboard/comment_review.html', {
                    'comment': comment
                })
            comment.status = 'not_approved'
            comment.justification = justification
            messages.success(request, "Comment disapproved successfully.")
        else:
            messages.error(request, "Invalid action.")
            return redirect('comments-list', factor_id=comment.factor.id)

        comment.save()
        return redirect('comments-list', factor_id=comment.factor.id)
    
    
@method_decorator(login_required, name="dispatch")
class JustificationDetailView(View):
    def get(self, request, comment_id):
        comment = get_object_or_404(CommentsModel, id=comment_id)
        if not comment.justification:
            messages.error(request, "No justification available for this comment.")
            return redirect('comments-list', factor_id=comment.factor.id)
        return render(request, 'dashboard/justification_detail.html', {
            'comment': comment
        })

# ... (otras importaciones y vistas existentes permanecen iguales)

@method_decorator(login_required, name="dispatch")
class CharacteristicManageView(View):
    def get(self, request, factor_id):
        factor = get_object_or_404(FactorModel, id=factor_id)
        characteristics = factor.characteristics.all()
        return render(request, 'dashboard/characteristic_manage.html', {
            'factor': factor,
            'characteristics': characteristics,
            'user_role': getattr(request.user, 'role', None)
        })

@method_decorator(login_required, name="dispatch")
class CharacteristicDevelopView(View):
    def get(self, request, factor_id, characteristic_id):
        factor = get_object_or_404(FactorModel, id=factor_id)
        characteristic = get_object_or_404(CharacteristicModel, id=characteristic_id)
        if characteristic.state.name == "Completed":
            messages.error(request, "Cannot develop a completed characteristic.")
            return redirect('characteristic-manage', factor_id=factor.id)
        form = CharacteristicDevelopForm()
        return render(request, 'dashboard/characteristic_develop.html', {
            'factor': factor,
            'characteristic': characteristic,
            'form': form
        })

    def post(self, request, factor_id, characteristic_id):
        factor = get_object_or_404(FactorModel, id=factor_id)
        characteristic = get_object_or_404(CharacteristicModel, id=characteristic_id)
        if characteristic.state.name == "Completed":
            messages.error(request, "Cannot develop a completed characteristic.")
            return redirect('characteristic-manage', factor_id=factor.id)
        form = CharacteristicDevelopForm(request.POST)
        if form.is_valid():
            strength_new = form.cleaned_data['strength_new']
            strength_existent = form.cleaned_data['strength_existent']
            aspect_new = form.cleaned_data['aspect_new']
            aspect_existent = form.cleaned_data['aspect_existent']

            if strength_new:
                strength, created = GlobalStrengths.objects.get_or_create(name=strength_new, created_by=request.user)
                CharacteristicStrengths.objects.get_or_create(characteristic=characteristic, global_strength=strength)
            if strength_existent:
                CharacteristicStrengths.objects.get_or_create(characteristic=characteristic, global_strength=strength_existent)
            if aspect_new:
                aspect, created = GlobalAspects.objects.get_or_create(name=aspect_new, created_by=request.user)
                CharacteristicAspects.objects.get_or_create(characteristic=characteristic, global_aspect=aspect)
            if aspect_existent:
                CharacteristicAspects.objects.get_or_create(characteristic=characteristic, global_aspect=aspect_existent)

            messages.success(request, "Strengths and aspects updated successfully.")
            return redirect('characteristic-develop', factor_id=factor.id, characteristic_id=characteristic.id)
        return render(request, 'dashboard/characteristic_develop.html', {
            'factor': factor,
            'characteristic': characteristic,
            'form': form
        })

@method_decorator(login_required, name="dispatch")
class CharacteristicCompleteView(View):
    def post(self, request, factor_id, characteristic_id):
        factor = get_object_or_404(FactorModel, id=factor_id)
        characteristic = get_object_or_404(CharacteristicModel, id=characteristic_id)
        if not characteristic.has_strengths_and_aspects:
            return JsonResponse({'status': 'error', 'message': 'Cannot complete without at least one strength and one aspect.'}, status=400)
        if getattr(request.user, 'role', None) and request.user.role.name != 'acadi':
            return JsonResponse({'status': 'error', 'message': 'Only "acadi" can complete a characteristic.'}, status=403)
        try:
            completed_state, created = States.objects.get_or_create(name="Completed")
            characteristic.state = completed_state
            characteristic.save()
            factor.save()  # Update factor progress
            messages.success(request, "Characteristic marked as completed.")  # Add user feedback
            return JsonResponse({'status': 'success', 'message': 'Characteristic marked as completed.'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': f'Error completing characteristic: {str(e)}'}, status=500)


@method_decorator(login_required, name="dispatch")
class CharacteristicCreateView(View):
    def get(self, request, factor_id):
        factor = get_object_or_404(FactorModel, id=factor_id)
        form = CharacteristicForm()
        return render(request, 'dashboard/characteristic_form.html', {'form': form, 'factor': factor})

    def post(self, request, factor_id):
        factor = get_object_or_404(FactorModel, id=factor_id)
        form = CharacteristicForm(request.POST)
        if form.is_valid():
            characteristic = form.save(commit=False)
            characteristic.created_by = request.user
            characteristic.save()
            characteristic.factors.add(factor)
            factor.save()  # Update factor progress
            messages.success(request, "Characteristic created successfully.")
            return redirect('characteristic-manage', factor_id=factor.id)
        return render(request, 'dashboard/characteristic_form.html', {'form': form, 'factor': factor})
    

@method_decorator(login_required, name="dispatch")
class CharacteristicUpdateView(View):
    def get(self, request, factor_id, characteristic_id):
        factor = get_object_or_404(FactorModel, id=factor_id)
        characteristic = get_object_or_404(CharacteristicModel, id=characteristic_id)
        form = CharacteristicForm(instance=characteristic)
        return render(request, 'dashboard/characteristic_form.html', {'form': form, 'factor': factor, 'characteristic': characteristic})

    def post(self, request, factor_id, characteristic_id):
        factor = get_object_or_404(FactorModel, id=factor_id)
        characteristic = get_object_or_404(CharacteristicModel, id=characteristic_id)
        form = CharacteristicForm(request.POST, instance=characteristic)
        if form.is_valid():
            form.save()
            messages.success(request, "Characteristic updated successfully.")
            return redirect('characteristic-manage', factor_id=factor.id)
        return render(request, 'dashboard/characteristic_form.html', {'form': form, 'factor': factor, 'characteristic': characteristic})

@method_decorator(login_required, name="dispatch")
class CharacteristicDeleteView(View):
    def dispatch(self, request, *args, **kwargs):
        user_role = getattr(request.user, 'role', None)
        if not user_role or user_role.name not in ['acadi', 'program director']:
            messages.error(request, "Only users with role 'acadi' or 'program director' can delete characteristics.")
            return redirect('characteristic-manage', factor_id=kwargs['factor_id'])
        return super().dispatch(request, *args, **kwargs)

    def post(self, request, factor_id, characteristic_id):
        factor = get_object_or_404(FactorModel, id=factor_id)
        characteristic = get_object_or_404(CharacteristicModel, id=characteristic_id)
        characteristic.delete()
        factor.save()  # Update factor progress
        messages.success(request, "Characteristic deleted successfully.")
        return redirect('characteristic-manage', factor_id=factor.id)

@method_decorator(login_required, name="dispatch")
class CharacteristicUploadCSVView(View):
    def dispatch(self, request, *args, **kwargs):
        user_role = getattr(request.user, 'role', None)
        if not user_role or user_role.name not in ['acadi', 'program director']:
            messages.error(request, "Only users with role 'acadi' or 'program director' can upload characteristics.")
            return redirect('characteristic-manage', factor_id=kwargs['factor_id'])
        return super().dispatch(request, *args, **kwargs)

    def get(self, request, factor_id):
        factor = get_object_or_404(FactorModel, id=factor_id)
        return render(request, 'dashboard/characteristic_upload.html', {'factor': factor})

    def post(self, request, factor_id, *args, **kwargs):
        factor = get_object_or_404(FactorModel, id=factor_id)
        if 'csv_file' in request.FILES:
            csv_file = request.FILES['csv_file']
            decoded_file = csv_file.read().decode('utf-8').splitlines()
            reader = csv.DictReader(decoded_file)
            characteristics_data = []
            for row in reader:
                if 'title' not in row:
                    messages.error(request, "CSV must contain a 'title' column.")
                    return redirect('characteristic-upload', factor_id=factor.id)
                characteristics_data.append({
                    'title': row['title'],
                    'description': row.get('description', '')
                })
            request.session['characteristics_data'] = characteristics_data
            request.session['factor_id'] = factor_id
            return render(request, 'dashboard/characteristic_preview.html', {
                'factor': factor,
                'characteristics_data': characteristics_data
            })
        messages.error(request, "Please upload a valid CSV file.")
        return redirect('characteristic-manage', factor_id=factor.id)

@method_decorator(login_required, name="dispatch")
class CharacteristicConfirmView(View):
    def get(self, request, factor_id):
        factor = get_object_or_404(FactorModel, id=factor_id)
        characteristics_data = request.session.get('characteristics_data', [])
        if not characteristics_data:
            messages.error(request, "No data to confirm.")
            return redirect('characteristic-manage', factor_id=factor.id)
        for data in characteristics_data:
            characteristic = CharacteristicModel(
                title=data['title'],
                description=data['description'],
                created_by=request.user
            )
            characteristic.save()
            characteristic.factors.add(factor)
        factor.save()  # Update factor progress
        del request.session['characteristics_data']
        del request.session['factor_id']
        messages.success(request, "Characteristics uploaded successfully.")
        return redirect('characteristic-manage', factor_id=factor.id)
    
@method_decorator(login_required, name="dispatch")
class CharacteristicDetailsView(View):
    def get(self, request, factor_id, characteristic_id):
        factor = get_object_or_404(FactorModel, id=factor_id)
        characteristic = get_object_or_404(CharacteristicModel, id=characteristic_id)
        if characteristic.state.name != "Completed":
            messages.error(request, "Details are only available for completed characteristics.")
            return redirect('characteristic-manage', factor_id=factor.id)
        return render(request, 'dashboard/characteristic_details.html', {
            'factor': factor,
            'characteristic': characteristic
        })
    
@login_required
def profile_view(request):
    return render(request, 'dashboard/profile.html')
  

@method_decorator(login_required, name="dispatch")
class FactorUpdateView(View):
    def get(self, request, factor_id):
        factor = get_object_or_404(FactorModel, pk=factor_id)

        if not request.user.role or request.user.role.name not in ['acadi', 'program director']:
            messages.error(request, "You don't have permissions")
            return redirect('report-list')

        if not factor.report:
            messages.error(request, "This factor is not associated with any report.")
            return redirect('report-list')

        form = FactorForm(instance=factor)

        return render(request, 'dashboard/edit-factor.html', {
            'form': form,
            'factor': factor
        })
    

    def post(self, request, factor_id):
        factor = get_object_or_404(FactorModel, pk=factor_id)

        if not request.user.role or request.user.role.name not in ['acadi', 'program director']:
            messages.error(request, "You don't have permissions")
            return redirect('view-report', pk=factor.report.id)

        form = FactorForm(request.POST, instance=factor)

        if form.is_valid():
            form.save()
            messages.success(request, f"Factor '{factor.name}' updated succesfully")
            return redirect('view-report', pk=factor.report.id)

        return render(request, 'dashboard/edit-factor.html', {
            'form': form,
            'factor': factor
        })

@login_required
def edit_profile_view(request):
    if request.method == 'POST':
        form = ProfileForm(request.POST, request.FILES, instance=request.user)
        if form.is_valid():
            form.save()
            messages.success(request, "Foto de perfil actualizada correctamente.")
            return redirect('profile-view')
    else:
        form = ProfileForm(instance=request.user)
    return render(request, 'dashboard/edit_profile.html', {'form': form})



@method_decorator(login_required, name="dispatch")
class QuestionManageView(View):
    def get(self, request, factor_id):
        factor = get_object_or_404(FactorModel, id=factor_id)
        questions = factor.questions.all()
        return render(request, 'dashboard/question_manage.html', {
            'factor': factor,
            'questions': questions,
            'is_acadi': request.user.role.name == 'acadi'
        })


@method_decorator(login_required, name="dispatch")
class QuestionCreateView(View):
    def dispatch(self, request, *args, **kwargs):
        user_role = getattr(request.user, 'role', None)
        if not user_role or user_role.name != 'acadi':
            messages.error(request, "Solo los usuarios con rol 'acadi' pueden crear preguntas.")
            return redirect('question-manage', factor_id=kwargs['factor_id'])
        return super().dispatch(request, *args, **kwargs)

    def get(self, request, factor_id):
        factor = get_object_or_404(FactorModel, id=factor_id)
        form = QuestionForm()
        return render(request, 'dashboard/question_form.html', {'form': form, 'factor': factor})

    def post(self, request, factor_id):
        factor = get_object_or_404(FactorModel, id=factor_id)
        form = QuestionForm(request.POST, request.FILES)
        if form.is_valid():
            question = form.save(commit=False)
            question.owner = request.user
            question.factor = factor
            txt_content = form.cleaned_data.get('txt_file')
            if txt_content:
                question.description = txt_content
            question.save()
            messages.success(request, "Pregunta creada exitosamente.")
            return redirect('question-manage', factor_id=factor.id)
        return render(request, 'dashboard/question_form.html', {'form': form, 'factor': factor})
    


@method_decorator(login_required, name="dispatch")
class TaskAssignView(View):
    def dispatch(self, request, *args, **kwargs):
        user_role = getattr(request.user, 'role', None)
        if not user_role or user_role.name not in ['acadi', 'program director']:
            messages.error(request, "Only users with role 'acadi' or 'program director' can assign tasks.")
            return redirect('dashboard-admin')
        return super().dispatch(request, *args, **kwargs)

    def get(self, request):
        form = TaskForm()
        users = UserModel.objects.all()
        return render(request, 'dashboard/task-assign.html', {'form': form, 'users': users})

    def post(self, request):
        form = TaskForm(request.POST)
        if form.is_valid():
            task = form.save(commit=False)
            task.created_by = request.user
            task.save()
            NotificationModel.objects.create(
                title=f"Task: {task.title} Due: {task.due_date}",
                user=task.assignee,
                created_by=request.user
            )
            messages.success(request, f"Task '{task.title}' assigned to {task.assignee.username} successfully.")
            return redirect('task-assign')
        users = UserModel.objects.all()
        return render(request, 'dashboard/task-assign.html', {'form': form, 'users': users})
    
@method_decorator(login_required, name="dispatch")
class FactorCollaborativeEditView(View):
    def get(self, request, factor_id):
        factor = get_object_or_404(FactorModel, id=factor_id)
        
        if not factor.google_doc_url:
            messages.error(request, "There isn't google docs in this factor")
            return redirect('view-report', pk=factor.report.id)
        
        return HttpResponseRedirect(factor.google_doc_url)

@method_decorator(login_required, name="dispatch")
class SendAccreditationNotificationView(View):
    def get(self, request):
        return render(request, 'dashboard/send_notification.html')

    def post(self, request):
        recipient_email = request.POST.get('recipient_email') 
        start_date = request.POST.get('start_date')

        if not recipient_email:
            messages.error(request, "must enter a recipient email.")
            return redirect('send_notification')

        try:
            validate_email(recipient_email)
        except ValidationError:
            messages.error(request, "Please enter a valid email address.")
            return redirect('send_notification')

        if not start_date:
            messages.error(request, "You must specify the start date of the process.")
            return redirect('send_notification')

        subject = 'Convocatoria al Proceso de Acreditación'
        message = f"Usted ha sido convocado a hacer parte del proceso de acreditaciones, que inicia el {start_date}."
        from_email = settings.DEFAULT_FROM_EMAIL

        try:
            send_mail(
                subject,
                message,
                from_email,
                [recipient_email],
                fail_silently=False,
            )
            NotificationLog.objects.create(
                recipient=recipient_email,
                created_by=request.user
            )
            messages.success(request, f"Notification sent successfully to {recipient_email}.")
        except Exception as e:
            messages.error(request, f"Error sending notification: {str(e)}")

        return redirect('send_notification')

@method_decorator(login_required, name="dispatch")
class NotificationHistoryView(View):
    def get(self, request):
        logs = NotificationLog.objects.all().order_by('-sent_at')
        return render(request, 'dashboard/notification_history.html', {'logs': logs})

@method_decorator(login_required, name="dispatch")
class StartAccreditationProcessView(View):
    def get(self, request):
        if request.user.role.name != 'acadi':  # Ajusta según tu modelo de roles
            messages.error(request, 'No tienes permiso para iniciar un proceso de acreditación.')
            return redirect('dashboard-admin')
        return render(request, 'dashboard/start_accreditation.html')

    def post(self, request):
        name = request.POST.get('name')
        start_date = request.POST.get('start_date')

        if not name:
            messages.error(request, 'Debes especificar un nombre para el proceso.')
            return redirect('start-accreditation')
        if not start_date:
            messages.error(request, 'Debes especificar la fecha de inicio del proceso.')
            return redirect('start-accreditation')

        try:
            AccreditationProcess.objects.create(
                name=name,
                start_date=start_date,
                status='planning',
                created_by=request.user
            )
            messages.success(request, 'Proceso de acreditación iniciado exitosamente.')
        except Exception as e:
            messages.error(request, f'Error al iniciar el proceso: {str(e)}')

        return redirect('dashboard-admin')
    
@method_decorator(login_required, name="dispatch")
class GetAccreditationStatusView(View):
    def get(self, request):
        try:
            process = AccreditationProcess.objects.latest('created_at')
            current_date = timezone.now().date()

            # Calculate days elapsed
            days_elapsed = (current_date - process.start_date).days if process.start_date else 0
            print(f"Current date: {current_date}")
            print(f"Start date: {process.start_date}")
            print(f"Days elapsed: {days_elapsed}")
            
            # Define stage durations (in months converted to days)
            stages = {
                'planning': 90,        # 3 months
                'analysis': 90,        # 3 months (parallel with planning)
                'evidence': 150,       # 5 months
                'report': 180,         # 6 months
                'consolidation': 120,  # 4 months
                'submission': 30       # 1 month
            }
            
            total_days = sum(stages.values())
            progress_percentage = min(100, (days_elapsed / total_days) * 100)
            
            print(f"Total days in process: {total_days}")
            print(f"Progress percentage: {progress_percentage}%")
            
            # Determine current stage based on days elapsed
            current_status = 'pending'
            accumulated_days = 0
            
            for stage, duration in stages.items():
                accumulated_days += duration
                print(f"Checking stage {stage}: {accumulated_days} days")
                if days_elapsed <= accumulated_days:
                    current_status = stage
                    print(f"Current stage found: {stage}")
                    break
                elif days_elapsed > total_days:
                    current_status = 'completed'
                    print("Process completed")
                    break

            return JsonResponse({
                'status': current_status,
                'start_date': process.start_date.strftime('%Y-%m-%d') if process.start_date else None,
                'days_elapsed': days_elapsed,
                'progress_percentage': progress_percentage
            })
        except AccreditationProcess.DoesNotExist:
            print("No accreditation process found")
            return JsonResponse({'status': 'pending', 'start_date': None})

        
@method_decorator(login_required, name='dispatch')
class DOFADocumentView(View):
    def get(self, request):
        # Get all factors with their characteristics
        factors = FactorModel.objects.all().prefetch_related(
            'characteristics',
            'characteristics__strengths__global_strength',
            'characteristics__aspects__global_aspect'
        )
        
        # Create document
        doc = docx.Document()
        doc.add_heading('DOFA Analysis', 0)

        # Iterate through each factor
        for factor in factors:
            doc.add_heading(f'Factor: {factor.name}', level=1)
            
            # Add strengths section
            doc.add_heading('Strengths:', level=2)
            for char in factor.characteristics.all():
                for strength in char.strengths.all():
                    p = doc.add_paragraph()
                    p.add_run(f'• {strength.global_strength.name}')
            
            # Add aspects to improve section
            doc.add_heading('Aspects to Improve:', level=2)
            for char in factor.characteristics.all():
                for aspect in char.aspects.all():
                    p = doc.add_paragraph()
                    p.add_run(f'• {aspect.global_aspect.name}')
            
            # Add page break between factors
            doc.add_page_break()

        # Create response with the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=DOFA_Analysis.docx'
        doc.save(response)
        
        return response
